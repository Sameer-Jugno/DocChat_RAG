"""Extract text from supported document formats into PageText units."""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

from pypdf import PdfReader

from app.config import get_settings
from app.ingest.text_utils import normalize_text
from app.ingest.types import PageText

logger = logging.getLogger("pdf_chat.loader")

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".txt", ".md", ".markdown", ".csv", ".docx"})

ACCEPT_MAP = {
    "application/pdf": [".pdf"],
    "application/x-pdf": [".pdf"],
    "text/plain": [".txt", ".md", ".markdown", ".csv"],
    "text/markdown": [".md", ".markdown"],
    "text/csv": [".csv"],
    "application/csv": [".csv"],
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
    "application/octet-stream": [".pdf", ".txt", ".md", ".csv", ".docx"],
}


def is_supported(path_or_name: str | Path) -> bool:
    return Path(path_or_name).suffix.lower() in SUPPORTED_EXTENSIONS


def supported_list() -> str:
    return ", ".join(sorted({ext for ext in SUPPORTED_EXTENSIONS}))


def extract_document(path: Path) -> list[PageText]:
    """Dispatch extraction by file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return _extract_plain_text(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(
        f"Unsupported file type `{suffix}`. Supported: {supported_list()}"
    )


def _extract_pdf(path: Path) -> list[PageText]:
    """
    Extract PDF text.

    1) Native text via pypdf
    2) If a page has little/no text and OCR is enabled → render page + OCR
    3) Optionally OCR embedded images on text pages and append
    """
    settings = get_settings()
    reader = PdfReader(str(path))
    native: list[str] = []
    for page in reader.pages:
        native.append(normalize_text(page.extract_text() or ""))

    need_ocr_pages = [
        i
        for i, text in enumerate(native)
        if len(text) < settings.ocr_min_chars
    ]

    ocr_page_text: dict[int, str] = {}
    ocr_image_text: dict[int, str] = {}

    if settings.ocr_enabled and (
        need_ocr_pages or settings.ocr_embedded_images
    ):
        try:
            ocr_page_text, ocr_image_text = _ocr_pdf_pages(
                path,
                page_indexes=need_ocr_pages if need_ocr_pages else [],
                also_images=settings.ocr_embedded_images,
                dpi=settings.ocr_dpi,
            )
        except Exception:  # noqa: BLE001
            logger.exception("PDF OCR pipeline failed; continuing with native text only")

    pages: list[PageText] = []
    for i, native_text in enumerate(native):
        parts: list[str] = []
        if len(native_text) >= settings.ocr_min_chars:
            parts.append(native_text)
        elif i in ocr_page_text and ocr_page_text[i]:
            parts.append(ocr_page_text[i])
        elif native_text:
            parts.append(native_text)

        img_extra = ocr_image_text.get(i, "")
        if img_extra:
            # Avoid duplicating if page OCR already captured the same content
            if not parts or img_extra not in parts[0]:
                parts.append(f"[Image text]\n{img_extra}")

        merged = normalize_text("\n\n".join(p for p in parts if p))
        if merged:
            pages.append(PageText(page_number=i + 1, text=merged, unit="page"))

    return pages


def _ocr_pdf_pages(
    path: Path,
    *,
    page_indexes: list[int],
    also_images: bool,
    dpi: int,
) -> tuple[dict[int, str], dict[int, str]]:
    """Return (page_ocr_by_index, embedded_image_ocr_by_index)."""
    import pymupdf as fitz

    from app.ingest.ocr import ocr_image, ocr_pixmap_png_bytes

    page_ocr: dict[int, str] = {}
    image_ocr: dict[int, str] = {}
    doc = fitz.open(str(path))
    zoom = max(dpi, 72) / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    page_set = set(page_indexes)

    try:
        for i in range(doc.page_count):
            page = doc.load_page(i)

            if i in page_set:
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                text = normalize_text(ocr_pixmap_png_bytes(pix.tobytes("png")))
                if text:
                    page_ocr[i] = text
                    logger.info("OCR page %s → %s chars", i + 1, len(text))

            if also_images and i not in page_set:
                texts: list[str] = []
                for img in page.get_images(full=True):
                    xref = img[0]
                    try:
                        extracted = doc.extract_image(xref)
                    except Exception:  # noqa: BLE001
                        continue
                    image_bytes = extracted.get("image")
                    if not image_bytes:
                        continue
                    # Skip tiny icons
                    w = int(extracted.get("width") or 0)
                    h = int(extracted.get("height") or 0)
                    if w * h < 80_000:
                        continue
                    try:
                        from io import BytesIO

                        from PIL import Image

                        pil = Image.open(BytesIO(image_bytes))
                        # Cap huge images for speed
                        max_side = 1600
                        if max(pil.size) > max_side:
                            pil.thumbnail((max_side, max_side))
                        t = normalize_text(ocr_image(pil))
                    except Exception:  # noqa: BLE001
                        logger.debug("embedded image OCR failed on page %s", i + 1)
                        continue
                    if t and len(t) >= 20:
                        texts.append(t)
                if texts:
                    image_ocr[i] = "\n".join(texts)
    finally:
        doc.close()

    return page_ocr, image_ocr


def _extract_plain_text(path: Path) -> list[PageText]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = normalize_text(raw)
    if not text:
        return []
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(parts) <= 1:
        return [PageText(page_number=1, text=text, unit="section")]

    sections: list[str] = []
    buf: list[str] = []
    size = 0
    target = 1500
    for part in parts:
        if size + len(part) > target and buf:
            sections.append("\n\n".join(buf))
            buf = [part]
            size = len(part)
        else:
            buf.append(part)
            size += len(part) + 2
    if buf:
        sections.append("\n\n".join(buf))

    return [
        PageText(page_number=i + 1, text=sec, unit="section")
        for i, sec in enumerate(sections)
    ]


def _extract_csv(path: Path, rows_per_unit: int = 25) -> list[PageText]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return []

    header = rows[0]
    header_line = " | ".join(header)
    data_rows = rows[1:] if len(rows) > 1 else []

    if not data_rows:
        text = normalize_text(header_line)
        return [PageText(page_number=1, text=text, unit="rows")] if text else []

    units: list[PageText] = []
    for i in range(0, len(data_rows), rows_per_unit):
        batch = data_rows[i : i + rows_per_unit]
        lines = [f"Columns: {header_line}", ""]
        for ridx, row in enumerate(batch, start=i + 1):
            cells = list(row) + [""] * max(0, len(header) - len(row))
            pairs = [f"{h}={c}" for h, c in zip(header, cells)]
            lines.append(f"Row {ridx}: " + "; ".join(pairs))
        text = normalize_text("\n".join(lines))
        if text:
            units.append(
                PageText(page_number=len(units) + 1, text=text, unit="rows")
            )
    return units


def _extract_docx(path: Path) -> list[PageText]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError(
            "python-docx is required for .docx files. pip install python-docx"
        ) from exc

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        return []

    sections: list[str] = []
    buf: list[str] = []
    size = 0
    target = 1500
    for para in paragraphs:
        if size + len(para) > target and buf:
            sections.append("\n\n".join(buf))
            buf = [para]
            size = len(para)
        else:
            buf.append(para)
            size += len(para) + 2
    if buf:
        sections.append("\n\n".join(buf))

    return [
        PageText(page_number=i + 1, text=normalize_text(sec), unit="section")
        for i, sec in enumerate(sections)
        if normalize_text(sec)
    ]
